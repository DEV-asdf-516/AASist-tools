from abc import ABC, abstractmethod
from enum import Enum, auto
import os
from typing import (
    Dict,
    Iterable,
    List,
    Tuple,
)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import pandas as pd
from gui.handler import _GUIDANCE_LOG_NAME, LogLevel, QueueHandler
from guidance.schema_types import TableFormat
from guidance.submodel_table_parser import ParseObject, SubmodelTableParser
from openpyxl import load_workbook


class RowPipelineStage(Enum):
    idle = auto()
    set_id_short = auto()
    set_model_value = auto()
    set_semantic_id = auto()
    set_description = auto()
    set_MLP_model_value = auto()
    flush = auto()


class ConceptDescriptionPipelineStage(Enum):
    idle = auto()
    set_definition = auto()
    flush = auto()


class SubmodelTableExtractor(ABC):

    def __init__(
        self,
        parser: SubmodelTableParser,
        submodels: List[str] = None,
        columns: List[str] = None,
        **kwargs,
    ):
        self._parser = parser
        self._submodel_store: Dict[str, ParseObject] = {}
        self.submodels = submodels
        self.columns = columns
        self._file_name = kwargs.get("file_name", None)

    def export(self, format: TableFormat, **kwargs):
        self._prefix = (self._file_name or "output").split(".")[0]

        log_handler: QueueHandler = kwargs.get(
            "log_handler", QueueHandler(_GUIDANCE_LOG_NAME)
        )
        try:
            for submodel, df in self._to_dataframes():
                if format == TableFormat.DOCX:
                    docx = Document()
                    table = docx.add_table(rows=1, cols=len(df.columns))
                    table.style = "Table Grid"

                    for i, column in enumerate(df.columns):
                        cell = table.cell(0, i)
                        cell.text = column
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        if i > 0:
                            left = table.cell(0, i - 1)
                            if not left.text:
                                left.merge(cell)
                                left.text = left.text.strip()
                                left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    for row in df.itertuples(index=False):
                        cells = table.add_row().cells
                        for i, value in enumerate(row):
                            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cells[i].text = (
                                ""
                                if value is None or not isinstance(value, str)
                                else str(value)
                            )
                    log_handler.add(
                        f"Exporting submodel '{submodel}' to {self._prefix}_{submodel}.docx"
                    )
                    docx.save(self._prefix + "_" + submodel + ".docx")

                if format == TableFormat.XLSX:
                    temp_file_name = self._prefix + "_" + submodel + "_temp.xlsx"
                    df.to_excel(temp_file_name, index=False)
                    wb = load_workbook(temp_file_name)
                    ws = wb.active

                    for i, cell in enumerate(ws[1], start=1):
                        if not cell.value and i < len(ws[1]) - 1:
                            right = ws.cell(row=1, column=i + 1)
                            if right.value:
                                cell.value = right.value
                                ws.merge_cells(
                                    start_row=1,
                                    start_column=i,
                                    end_row=1,
                                    end_column=i + 1,
                                )

                    log_handler.add(
                        f"Exporting submodel '{submodel}' to {self._prefix}_{submodel}.xlsx"
                    )
                    wb.save(self._prefix + "_" + submodel + ".xlsx")
                    wb.close()
                    os.remove(temp_file_name)
        except Exception as e:
            log_handler.add(
                f"Error exporting submodel '{submodel}' : {e.__class__}",
                log_level=LogLevel.ERROR,
            )

    @abstractmethod
    def extract_table(self):
        pass

    @abstractmethod
    def _to_dataframes(self) -> Iterable[Tuple[str, pd.DataFrame]]:
        pass
